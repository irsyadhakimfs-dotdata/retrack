# Generator file .docx dari template PPT ReTrack DWH.
# Jalankan: python dwh/build_docx.py  ->  menghasilkan dwh/PPT-ReTrack-DWH.docx
# Tiap "slide" jadi satu seksi: judul (Heading 1) + bullet isi + narasi + screenshot.

from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

TEAL = RGBColor(0x0A, 0x4A, 0x4A)
TEAL_BRAND = RGBColor(0x33, 0x99, 0x99)
PINK = RGBColor(0xC4, 0x44, 0x66)
GRAY = RGBColor(0x55, 0x55, 0x55)
CODE_BG = "EAF6F6"


def shade(paragraph, fill):
    """Beri warna latar (shading) pada paragraf — untuk blok kode."""
    pPr = paragraph._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), fill)
    pPr.append(shd)


def add_code(doc, code):
    """Tambah blok kode monospace dengan latar teal muda."""
    for line in code.strip("\n").split("\n"):
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Inches(0.3)
        p.paragraph_format.space_after = Pt(0)
        p.paragraph_format.space_before = Pt(0)
        shade(p, CODE_BG)
        run = p.add_run(line if line else " ")
        run.font.name = "Consolas"
        run.font.size = Pt(9)
        rPr = run._element.get_or_add_rPr()
        rFonts = rPr.get_or_add_rFonts()
        rFonts.set(qn("w:ascii"), "Consolas")
        rFonts.set(qn("w:hAnsi"), "Consolas")


def add_slide(doc, title):
    """Mulai seksi slide baru dengan judul Heading 1."""
    h = doc.add_heading(title, level=1)
    for run in h.runs:
        run.font.color.rgb = TEAL


def add_label(doc, text, color):
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.bold = True
    r.font.color.rgb = color
    r.font.size = Pt(10)
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(2)


def bullets(doc, items):
    for it in items:
        p = doc.add_paragraph(style="List Bullet")
        p.add_run(it)


def narasi(doc, text):
    p = doc.add_paragraph()
    r = p.add_run("🎤 Narasi: ")
    r.bold = True
    r.font.color.rgb = TEAL_BRAND
    r2 = p.add_run(text)
    r2.italic = True
    r2.font.color.rgb = GRAY


def shot(doc, text):
    p = doc.add_paragraph()
    r = p.add_run("📷 Screenshot: ")
    r.bold = True
    r.font.color.rgb = PINK
    r2 = p.add_run(text)
    r2.font.color.rgb = GRAY


doc = Document()

# Gaya default Normal -> Manrope-ish fallback Calibri, ukuran nyaman
style = doc.styles["Normal"]
style.font.name = "Calibri"
style.font.size = Pt(11)

# ---------- COVER ----------
t = doc.add_paragraph()
t.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = t.add_run("ReTrack")
r.bold = True
r.font.size = Pt(34)
r.font.color.rgb = TEAL_BRAND

st = doc.add_paragraph()
st.alignment = WD_ALIGN_PARAGRAPH.CENTER
rs = st.add_run("Perancangan Data Warehouse untuk Personal Finance Tracker")
rs.font.size = Pt(16)
rs.font.color.rgb = TEAL

for line in [
    "Tugas Project Based Data Warehouse Ke-II — T.A. 2025/2026",
    "Nama / NIM : ____________________",
    "Program Studi Data Science",
    "Dosen Pengampu : ____________________",
    "Deadline PPT: Minggu, 31 Mei 2026 · Presentasi: Rabu, 3 Juni 2026",
]:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run(line).font.size = Pt(11)

doc.add_paragraph()
note = doc.add_paragraph()
note.alignment = WD_ALIGN_PARAGRAPH.CENTER
rn = note.add_run(
    "Catatan: tiap 'Slide N' = satu slide PPT. Bagian Narasi = yang diucapkan "
    "(jangan ditulis di slide). Bagian Screenshot = gambar yang harus disiapkan."
)
rn.italic = True
rn.font.size = Pt(9)
rn.font.color.rgb = GRAY

doc.add_page_break()

# ---------- DATA SLIDE ----------
SQL_DDL_DATE = """CREATE TABLE dwh_dim_date (
    date_id     VARCHAR(8) PRIMARY KEY,   -- "20260531"
    full_date   DATE NOT NULL,
    day         INTEGER,
    month       INTEGER,
    month_name  VARCHAR(20),              -- "Mei"
    quarter     INTEGER,
    year        INTEGER,
    day_of_week VARCHAR(10),              -- "Sabtu"
    is_weekend  BOOLEAN
);"""

SQL_DDL_FACT = """CREATE TABLE dwh_fact_transaction (
    fact_id          INTEGER PRIMARY KEY,   -- = id transaksi OLTP (idempoten)
    date_id          VARCHAR(8) REFERENCES dwh_dim_date(date_id),
    user_id          INTEGER    REFERENCES dwh_dim_user(user_id),
    wallet_id        INTEGER    REFERENCES dwh_dim_wallet(wallet_id),
    category_id      INTEGER    REFERENCES dwh_dim_category(category_id),
    amount           NUMERIC(15,2),
    kind             VARCHAR(10),           -- income / expense
    usd_rate_at_date FLOAT,
    erosi_persen     FLOAT,
    etl_loaded_at    DATETIME
);"""

SQL_EXTRACT = """SELECT t.id, t.user_id, t.wallet_id, t.category_id,
       t.amount, t.kind, t.date, t.usd_rate_at_date
FROM transactions t
WHERE t.user_id = :user_id;"""

SQL_EROSI = """nilai_usd_awal     = amount / kurs_saat_transaksi
nilai_usd_sekarang = amount / kurs_sekarang
erosi_persen       = (nilai_usd_awal - nilai_usd_sekarang) / nilai_usd_awal * 100

Contoh: Rp5.000.000, kurs 15.000 -> 16.000  =>  erosi ~ 6,25%"""

SQL_LOAD = """INSERT OR REPLACE INTO dwh_fact_transaction
    (fact_id, date_id, user_id, wallet_id, category_id,
     amount, kind, usd_rate_at_date, erosi_persen, etl_loaded_at)
VALUES (:fact_id, :date_id, :user_id, :wallet_id, :category_id,
        :amount, :kind, :usd_rate, :erosi, CURRENT_TIMESTAMP);"""

SQL_SUMMARY = """SELECT d.month, d.month_name, f.kind, SUM(f.amount) AS total
FROM dwh_fact_transaction f
JOIN dwh_dim_date d ON f.date_id = d.date_id
WHERE d.year = 2026
GROUP BY d.month, d.month_name, f.kind
ORDER BY d.month;"""

# Tiap entri: (judul, fungsi pembangun isi)
add_slide(doc, "Slide 2 — Agenda / Outline")
bullets(doc, [
    "Latar belakang & topik (ReTrack)",
    "Pentingnya Data Warehouse — manfaat & tantangan",
    "Arsitektur DWH (Bottom-Up / Kimball)",
    "Desain: Conceptual -> Logical -> Physical",
    "Software perancangan DWH",
    "Proses ETL (Extract – Transform – Load)",
    "Software & hasil visualisasi data (dashboard)",
    "Kesimpulan",
])
narasi(doc, "Sebutkan agenda mengikuti 6 poin wajib di lembar tugas + intro & penutup.")

add_slide(doc, "Slide 3 — Latar Belakang Topik (ReTrack)")
bullets(doc, [
    "ReTrack = aplikasi web pencatat keuangan & budget planner untuk anak muda Indonesia (18–30 th).",
    "Sumber data operasional (OLTP): transaksi, dompet (wallet), kategori, budget, savings goal.",
    "Fitur khas: deteksi erosi nilai — seberapa tergerus daya beli pemasukan Rupiah akibat pelemahan kurs USD/IDR.",
    "Data transaksi terus bertumbuh -> butuh analisis multi-periode (bulanan, kuartalan, tahunan).",
])
narasi(doc, "ReTrack sudah punya database operasional; kita menambah lapisan DWH di atasnya untuk kebutuhan analitik.")
shot(doc, "(opsional) halaman Dashboard / Transaksi aplikasi ReTrack.")

add_slide(doc, "Slide 4 — Pentingnya Data Warehouse (Permasalahan)")
bullets(doc, [
    "Database OLTP dioptimalkan untuk transaksi (insert/update cepat), BUKAN agregasi besar.",
    "Query analitik (total pengeluaran per bulan, tren 12 bulan, top kategori) memberatkan tabel OLTP dan memperlambat aplikasi.",
    "Struktur OLTP ternormalisasi -> butuh banyak JOIN untuk laporan -> lambat & rumit.",
    "Tidak ada dimensi waktu siap pakai (quarter, nama bulan, weekend) untuk slicing data.",
])
narasi(doc, "Tekankan pemisahan beban: OLTP untuk operasional, DWH untuk analitik.")

add_slide(doc, "Slide 5 — Manfaat & Tantangan DWH")
add_label(doc, "Manfaat utama:", TEAL_BRAND)
bullets(doc, [
    "Query analitik cepat (skema denormalisasi, sedikit JOIN).",
    "Dimensi waktu siap pakai -> analisis tren, musiman, weekend vs weekday.",
    "Single source of truth untuk pelaporan & dashboard.",
    "Mendukung keputusan keuangan personal berbasis data.",
])
add_label(doc, "Tantangan implementasi:", PINK)
bullets(doc, [
    "Menjaga konsistensi data OLTP -> DWH (proses ETL harus benar).",
    "Idempotensi ETL — menjalankan ETL berulang tidak boleh menduplikasi data.",
    "Sinkronisasi data eksternal (kurs USD/IDR) untuk menghitung erosi.",
    "Penyimpanan ganda (data tersalin) & jadwal refresh data.",
])
narasi(doc, "Jujur sebutkan tantangan; tunjukkan solusi idempotensi: PK = id transaksi.")

add_slide(doc, "Slide 6 — Arsitektur DWH yang Digunakan")
bullets(doc, [
    "Pendekatan: Bottom-Up (Metodologi Kimball).",
    "Mulai dari satu data mart paling kritis: data mart Transaksi Keuangan.",
    "Bisa dikembangkan ke fakta lain (budget, goals) -> enterprise DWH bertahap.",
    "DWH disimpan satu basis data dengan OLTP, tabel dipisah dengan prefix dwh_.",
])
add_label(doc, "Kenapa Bottom-Up (bukan Top-Down/Inmon)?", TEAL_BRAND)
bullets(doc, [
    "Skala personal finance kecil -> tidak perlu enterprise DWH besar di awal.",
    "Cepat memberi nilai (dashboard langsung jadi), iteratif, biaya rendah.",
])
narasi(doc, "Top-Down (Inmon) bangun EDW dulu baru data mart; Bottom-Up (Kimball) bangun data mart dulu. Kita pilih Bottom-Up.")
shot(doc, "Diagram arsitektur: OLTP -> ETL -> DWH Star Schema -> Chart.js Dashboard.")

add_slide(doc, "Slide 7 — Desain (a) Conceptual Design")
add_label(doc, "Tabel Fakta — fact_transaction (grain: 1 baris per transaksi)", PINK)
bullets(doc, [
    "Measures: amount (SUM/AVG/COUNT), erosi_persen (AVG).",
    "Derived: total_income, total_expense, selisih (net).",
    "FK: date_id, user_id, wallet_id, category_id.",
])
add_label(doc, "Tabel Dimensi:", TEAL_BRAND)
bullets(doc, [
    "dim_date — day, month, month_name, quarter, year, day_of_week, is_weekend.",
    "dim_user — name, email, created_at.",
    "dim_wallet — name, type.",
    "dim_category — name, kind (income/expense).",
])
narasi(doc, "Jelaskan grain (level detail) fakta = per transaksi, lalu measure & dimensi pengirisnya.")

add_slide(doc, "Slide 8 — Desain (b) Logical Design: STAR SCHEMA")
add_code(doc, """          dim_date
              |
dim_category — fact_transaction — dim_wallet
              |
          dim_user""")
add_label(doc, "Kenapa Star Schema (bukan Snowflake / Galaxy)?", TEAL_BRAND)
bullets(doc, [
    "Tiap dimensi hanya 1 level hierarki -> tidak perlu dinormalisasi jadi snowflake.",
    "Lebih sedikit JOIN -> query agregasi lebih cepat.",
    "Mudah dipahami & dijelaskan (cocok presentasi & skala kecil).",
    "Belum perlu Galaxy (multi-fakta berbagi dimensi) karena baru 1 tabel fakta.",
])
narasi(doc, "Tunjukkan bentuk bintang: fakta di tengah, dimensi mengelilingi. Tegaskan alasan memilihnya.")
shot(doc, "Diagram dari file dwh/star-schema-mermaid.md (render di mermaid.live -> PNG), atau diagram relasi DB Browser.")

add_slide(doc, "Slide 9 — Desain (c) Physical Design — DDL Tabel Dimensi")
add_code(doc, SQL_DDL_DATE)
narasi(doc, "Sebut tipe data & kunci. Tampilkan juga dim_category & dim_wallet bila muat.")
shot(doc, "WAJIB — eksekusi CREATE TABLE di DB Browser + tabel muncul di panel struktur (output).")

add_slide(doc, "Slide 10 — Physical Design — DDL Tabel Fakta")
add_code(doc, SQL_DDL_FACT)
narasi(doc, "fact_id = PK dari id transaksi OLTP -> kunci idempotensi ETL. Semua FK menunjuk dimensi (bentuk star).")
shot(doc, "WAJIB — hasil eksekusi DDL + tabel dwh_fact_transaction di DB.")

add_slide(doc, "Slide 11 — Software Perancangan DWH")
add_label(doc, "Tools:", TEAL_BRAND)
bullets(doc, [
    "SQLite — engine basis data DWH (embedded, satu file .db).",
    "SQLAlchemy ORM + Flask-Migrate (Alembic) — definisi tabel & migrasi schema.",
    "DB Browser for SQLite — desain visual, jalankan & screenshot query.",
    "Python 3.11 + Flask sebagai aplikasi induk.",
])
add_label(doc, "Keunggulan:", TEAL_BRAND)
bullets(doc, [
    "Zero-config, gratis, ringan; satu file mudah dibawa & di-backup.",
    "Terintegrasi langsung dengan aplikasi (DWH satu DB dengan OLTP).",
])
add_label(doc, "Keterbatasan:", PINK)
bullets(doc, [
    "Bukan untuk skala enterprise / concurrent write tinggi.",
    "Tidak ada fitur OLAP bawaan (partisi, materialized view) seperti BigQuery/Snowflake.",
    "Tipe data terbatas (boolean & numeric disimpan sederhana).",
])
narasi(doc, "Jujur akui keterbatasan, tapi tepat untuk skala personal finance.")

add_slide(doc, "Slide 12 — Proses ETL — Gambaran Umum")
add_code(doc, """OLTP (transactions, users, wallets, categories)
        |  EXTRACT  (baca data sumber)
        v
   TRANSFORM (bangun dim_date, normalisasi dimensi, hitung erosi_persen)
        v
     LOAD (db.session.merge -> tabel dwh_*, idempoten)
        v
   DWH Star Schema  ->  return {loaded, skipped, errors}""")
bullets(doc, [
    "Dijalankan via endpoint POST /api/dwh/etl/run.",
    "ETL tidak pernah mengubah data OLTP (read-only terhadap sumber).",
    "Idempoten: dijalankan 2x tidak menambah baris (cek fact_id sudah ada -> skip).",
])
narasi(doc, "Jelaskan 3 tahap E-T-L singkat, lalu detail di slide berikut.")

add_slide(doc, "Slide 13 — ETL: EXTRACT")
bullets(doc, [
    "Baca seluruh transaksi milik user dari tabel OLTP transactions.",
    "Baca data referensi: users, wallets, categories.",
    "Ambil kurs USD/IDR terkini (untuk erosi) — boleh gagal (fallback None).",
])
add_code(doc, SQL_EXTRACT)
shot(doc, "Hasil SELECT di atas (beberapa baris transaksi mentah).")

add_slide(doc, "Slide 14 — ETL: TRANSFORM")
bullets(doc, [
    "Generate dim_date dari tiap tanggal unik: date_id='YYYYMMDD', month_name & day_of_week Bahasa Indonesia, quarter, is_weekend.",
    "Normalisasi dimensi (upsert user/wallet/category — tanpa duplikat).",
    "Hitung erosi_persen memakai rumus daya beli:",
])
add_code(doc, SQL_EROSI)
narasi(doc, "Ini nilai tambah analitik kita — fakta diperkaya kolom turunan erosi_persen.")
shot(doc, "Contoh perhitungan (tabel kecil sebelum -> sesudah transform).")

add_slide(doc, "Slide 15 — ETL: LOAD")
bullets(doc, [
    "Muat ke tabel dwh_* memakai UPSERT (db.session.merge / INSERT OR REPLACE).",
    "fact_id = transaction.id -> idempoten, tak ada baris ganda.",
    "Catat waktu muat etl_loaded_at.",
    'Kembalikan statistik: {"loaded": N, "skipped": M, "errors": 0}.',
])
add_code(doc, SQL_LOAD)
shot(doc, 'WAJIB — respons JSON POST /api/dwh/etl/run -> {"ok":true,"data":{"loaded":..,"skipped":..,"errors":0}}.')
shot(doc, "Bukti idempoten: jalankan ETL 2x, kedua kalinya loaded:0, skipped:N.")

add_slide(doc, "Slide 16 — Software Visualisasi & Digital Dashboard")
bullets(doc, [
    "Software visualisasi: Chart.js (halaman /dwh-dashboard aplikasi ReTrack).",
    "Data dari endpoint query DWH: /api/dwh/summary, /api/dwh/top-categories, /api/dwh/erosi-summary.",
])
add_label(doc, "Yang divisualisasikan + jenis diagram:", TEAL_BRAND)
bullets(doc, [
    "Pemasukan vs Pengeluaran per bulan  (sumber: /summary)  -> Bar chart (grouped).",
    "Top 5 kategori pengeluaran  (sumber: /top-categories)  -> Doughnut / Pie chart.",
    "Rata-rata erosi nilai pemasukan  (sumber: /erosi-summary)  -> KPI Card (angka).",
])
narasi(doc, 'Contoh kalimat tugas: "Total pengeluaran berdasarkan kategori menggunakan diagram doughnut."')

add_slide(doc, "Slide 17 — Query Analitik DWH + Output")
add_code(doc, SQL_SUMMARY)
shot(doc, "WAJIB — hasil query di DB Browser DAN hasil chart di halaman DWH Dashboard.")
narasi(doc, "Query agregasi langsung memberi data untuk chart — inilah nilai DWH.")

add_slide(doc, "Slide 18 — Demo Dashboard (Screenshot)")
add_label(doc, "Screenshot WAJIB (halaman /dwh-dashboard):", PINK)
bullets(doc, [
    "Bar chart pemasukan vs pengeluaran per bulan.",
    "Doughnut top kategori pengeluaran.",
    "Card rata-rata erosi nilai + jumlah transaksi income.",
    'Label "Data Warehouse Mode".',
])
narasi(doc, "Demokan tombol 'Jalankan ETL' lalu chart ter-update; tutup dengan alur OLTP->ETL->DWH->Dashboard.")

add_slide(doc, "Slide 19 — Kesimpulan")
bullets(doc, [
    "DWH ReTrack memisahkan beban analitik dari operasional -> laporan cepat & ringan.",
    "Desain Star Schema (1 fakta + 4 dimensi) tepat untuk skala personal finance.",
    "Arsitektur Bottom-Up (Kimball) memberi data mart yang langsung berguna.",
    "ETL idempoten + kolom turunan erosi_persen = nilai analitik tambahan.",
    "Visualisasi Chart.js menjadikan data actionable bagi pengguna.",
])
add_label(doc, "Pengembangan ke depan:", TEAL_BRAND)
bullets(doc, [
    "Tambah fakta budget & goals (menuju Galaxy schema).",
    "Penjadwalan ETL otomatis.",
])

add_slide(doc, "Slide 20 — Penutup / Tanya Jawab")
bullets(doc, [
    "Terima kasih.",
    '"Ada pertanyaan?"',
    "Cantumkan kontak / nama tim.",
])

# ---------- CHECKLIST ----------
doc.add_page_break()
add_slide(doc, "Checklist Screenshot Wajib (sesuai tuntutan tugas)")
bullets(doc, [
    "Physical design — CREATE TABLE dimensi & fakta + tabel muncul di DB (Slide 9–10).",
    "ETL — SELECT extract, perhitungan transform, INSERT/UPSERT load + respons JSON etl/run (Slide 13–15).",
    "Visualisasi — query agregasi + chart di dashboard (Slide 17–18).",
])
p = doc.add_paragraph()
r = p.add_run("File SQL ilustrasi lengkap: dwh/dwh_star_schema.sql — buka di DB Browser for SQLite untuk eksekusi & screenshot.")
r.italic = True
r.font.color.rgb = GRAY

import os
out = os.path.join(os.path.dirname(__file__), "PPT-ReTrack-DWH.docx")
doc.save(out)
print("Tersimpan:", out)
